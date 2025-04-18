from docx import Document

def create_lesson_doc(user_name, standard):
    doc = Document()
    doc.add_heading(f"Lesson Plan for {user_name}", 0)
    doc.add_paragraph(f"Standard: {standard}")
    doc.add_paragraph("Generated lesson plan content goes here...")
    file_path = "lesson_plan.docx"
    doc.save(file_path)
    return file_path
